from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docs():
    doc = Document()

    # Title
    title = doc.add_heading('Australian Nodal Gas Market Optimization Model', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Strategic Multi-Year Dispatch and Infrastructure Planning Framework (2025–2050)')

    # Introduction
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This model is a mathematical optimization framework designed to simulate the Australian east coast gas market\'s '
        'transition through to 2050. It utilizes Nodal Least-Cost Dispatch logic to solve for the most efficient gas '
        'production, pipeline transmission, and infrastructure expansion decisions, while accounting for the structural '
        'shifts identified in the AEMO Gas Statement of Opportunities (GSOO).'
    )

    # Background
    doc.add_heading('2. Model Background & Objectives', level=1)
    doc.add_paragraph(
        'As Australia transitions towards Net Zero, the gas market faces a complex "dual challenge":'
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Basin Depletion: ').bold = True
    p.add_run('The rapid decline of traditional supply sources, particularly the Gippsland Basin in the Bass Strait.')
    p.add_row = doc.add_paragraph(style='List Bullet')
    p.add_row.add_run('Demand Transition: ').bold = True
    p.add_row.add_run('Structural demand reduction via electrification of residential heating, offset by persistent industrial needs and export commitments.')
    
    doc.add_paragraph(
        'The objective of the model is to minimize total system cost, which is the sum of production costs, '
        'transportation tariffs, capital expenditure for new builds, and the economic cost of unserved demand.'
    )

    # Data Sources
    doc.add_heading('3. Key Data Sources & Assumptions', level=1)
    doc.add_paragraph(
        'The model is grounded in data from the Australian Energy Market Operator (AEMO) 2026 forecasting cycle.'
    )

    doc.add_heading('3.1 Demand Traces (Selectable GSOO Baseline)', level=2)
    doc.add_paragraph(
        'The model demand is built from the AEMO 2026 GSOO. The user selects one of the three 2026 GSOO demand '
        'scenarios as the model baseline via a dropdown in the dashboard: Step Change (central, most likely path to '
        'Net Zero), Accelerated Transition (faster electrification - steepest residential/commercial decline, highest '
        'green-industrial growth, fastest LNG run-down), or Slower Growth (demand held higher for longer). Rather than '
        'applying arbitrary per-node growth rates, each demand sector is re-based directly onto its GSOO annual '
        'trajectory for the chosen baseline while the empirical daily profile shape from Gas Bulletin Board (GBB) '
        'actuals is preserved (shape correlation 1.0; only the annual level is scaled). GSOO indices are applied '
        'relative to 2026 and clamped to the 2026-2045 GSOO horizon, then held flat to 2050. The Winter, LNG and ADGSM '
        'scenario levers layer multiplicatively on top of whichever baseline is chosen. Queensland LNG export demand is '
        'calibrated to ~3,650 TJ/day (~1,250 PJ/year) in 2026. ("Progressive Change" is not a 2026 GSOO scenario - it '
        'belongs to the older 2025 GSOO vintage - so the model uses the consistent 2026 GSOO trio above.)'
    )
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sector / Node'
    hdr_cells[1].text = 'GSOO Driver (per chosen baseline)'
    hdr_cells[2].text = 'Rationale'

    data = [
        ['City nodes (Sydney, Melbourne, Adelaide, Brisbane)', 'GSOO ResComm trajectory (Fig. 17)', 'City-gate distribution demand; electrification-driven decline'],
        ['Gas-powered generation (GPG)', 'GSOO NEM trajectory + regional peaks', 'Annual level anchored to NEM GPG; winter-peaking from regional summer/winter peaks. AEMO publishes one 2026 GSOO GPG annual trajectory, so non-Step-Change baselines scale it by their regional GPG peak ratio vs Step Change'],
        ['Large industrial', 'GSOO industrial trajectory (year index)', 'Indexed on empirical BBLARGE levels (a subset of whole-sector industrial); Yarwun reclassified GPG -> industrial'],
        ['LNG cluster (APLNG, GLNG, QCLNG)', 'GSOO LNG trajectory (Fig. 19)', '~3,650 TJ/day in 2026; replaces previous flat assumption'],
    ]
    for region, rate, reason in data:
        row_cells = table.add_row().cells
        row_cells[0].text = region
        row_cells[1].text = rate
        row_cells[2].text = reason

    doc.add_paragraph(
        '\nDerived demand data is rebuilt from source by the "Regenerate All Data" button (src/regenerate_data.py), '
        'which runs the build pipeline in dependency order: build_curtailable_demand -> build_gsoo_scenarios -> '
        'build_gpg_demand_gsoo -> build_industrial_demand_gsoo -> build_demand_gsoo. The GSOO extract pulls all three '
        'baselines, and the demand builders emit one set of files per baseline (demand_<baseline>.csv etc.). Only '
        'source inputs (GBB actuals, GSOO workbooks, configs) are committed; generated files are rebuilt locally.'
    )

    doc.add_paragraph('\nRelevant Links:')
    doc.add_paragraph('AEMO 2026 GSOO Report Data:', style='List Bullet')
    doc.add_paragraph('https://www.aemo.com.au/-/media/files/gas/national_planning_and_forecasting/gsoo/2026/2026-gas-statement-of-opportunities-report-figures-and-data.xlsx', style='Caption')
    doc.add_paragraph('AEMO 2026 GSOO Supply Data:', style='List Bullet')
    doc.add_paragraph('https://www.aemo.com.au/-/media/files/gas/national_planning_and_forecasting/gsoo/2026/2026-gas-statement-of-opportunities-supply-data.xlsx', style='Caption')

    doc.add_heading('3.2 Supply Basin Dynamics', level=2)
    doc.add_paragraph(
        'Production capacities reflect the 2026 GSOO benchmarks. Decline rates for southern basins have been '
        'accelerated to reflect recent depletion reports.'
    )
    s_table = doc.add_table(rows=1, cols=3)
    s_table.rows[0].cells[0].text = 'Basin'
    s_table.rows[0].cells[1].text = 'Capacity (TJ/d)'
    s_table.rows[0].cells[2].text = 'Annual Decline Rate'
    s_data = [
        ['Gippsland (Bass Strait)', '766', '12.0%'],
        ['Moomba (Cooper Basin)', '400', '3.0%'],
        ['Surat (QLD CSG)', '4,000', '1.0%']
    ]
    for basin, cap, decline in s_data:
        r = s_table.add_row().cells
        r[0].text = basin
        r[1].text = cap
        r[2].text = decline

    # Technical implementation
    doc.add_heading('4. Technical Implementation', level=1)
    
    doc.add_heading('4.1 Mathematical Optimization', level=2)
    doc.add_paragraph(
        'The model uses Mixed-Integer Linear Programming (MILP), solved with the open-source HiGHS solver via '
        'Pyomo (appsi_highs). The core constraints include nodal mass balance (Supply + Inflow = Demand + Outflow), '
        'pipeline capacities, and storage continuity. Because HiGHS only returns dual values for a pure linear '
        'program, the binary build decisions are fixed and relaxed to continuous after the MILP step, and the model '
        'is re-solved as an LP to recover the shadow prices used for nodal price discovery.'
    )

    doc.add_heading('4.2 Storage Modeling', level=2)
    doc.add_paragraph(
        'Gas storage facilities (Iona, Moomba, Silver Springs) are modeled as "stateful" nodes. The inventory '
        'at the end of day (t) must equal the inventory from day (t-1) plus injections minus withdrawals. This '
        'allows the model to "pre-fill" storage in summer to meet the winter peaks triggered by the Southern Winter Stress scenarios.'
    )

    doc.add_heading('4.3 Multi-Year Path Dependency', level=2)
    doc.add_paragraph(
        'A unique feature of this simulation is the "already_built" logic. Infrastructure decisions are binary (0 or 1). '
        'Once the optimizer decides that a project like the Port Kembla LNG Import Terminal is economically necessary, '
        'that decision is locked in for all subsequent years in the 2050 sequence.'
    )

    # Pricing
    doc.add_heading('5. Price Discovery', level=1)
    doc.add_paragraph(
        'Nodal prices are derived from the marginal cost of supply (shadow prices). When a pipeline is uncongested, '
        'nodal prices are coupled (differing only by transport costs). When a pipeline hits its capacity limit, the '
        'nodal prices "decouple," reflecting local scarcity or surplus.'
    )

    # -----------------------------------------------------------------
    doc.add_heading('6. Curtailable Large-User Demand: GPG & Large Industrial', level=1)
    doc.add_paragraph(
        'Gas-powered generation (GPG) and large transmission-connected industrial users are '
        'modelled as explicit, curtailable demand tiers added on top of the distribution-level '
        'node demand. This was added in 2026 to capture demand that the nodal traces did not '
        'previously represent.'
    )

    doc.add_heading('6.1 Why these loads are additive (not double-counted)', level=2)
    doc.add_paragraph(
        'The model\'s nodal demand was calibrated from Gas Bulletin Board (GBB) Actual Flow data '
        'at the city distribution level. On the GBB, GPG plants (FacilityType BBGPG) and large '
        'industrial users (BBLARGE) are metered as their own facilities, separate from the '
        'distribution PIPE deliveries that the node traces track. A reconstruction of the demand '
        'decomposition from the same GBB source confirmed that node demand closely matches '
        'city-gate distribution delivery in every region (e.g. Sydney 235 vs 217 TJ/d; Adelaide '
        '53 vs 48; Brisbane 81 vs 63), with a worst-case residual overlap of ~2 TJ/d at Sydney. '
        'This is consistent with AEMO\'s own demand segmentation, which forecasts Tariff V '
        '(residential/commercial), Tariff D (industrial) and GPG as separate categories. GPG and '
        'large industrial are therefore genuinely additive, not embedded.'
    )

    doc.add_heading('6.2 Data source and seasonality', level=2)
    doc.add_paragraph(
        'Per-facility daily demand traces are derived from the GBB Actual Flow & Storage data '
        '(2022-2026), averaged into a representative 365-day shape per facility - the same method '
        'used for the distribution demand traces. Facilities are mapped to the nearest modelled '
        'node (NSW->Sydney; SA->Adelaide; VIC->Melbourne/Gippsland; QLD Surat basin->Surat, '
        'Gladstone cogen->Gladstone, SEQ->Brisbane). NT, Tasmania and off-network North '
        'Queensland (Mount Isa) facilities are excluded. Mapped totals are ~106 TJ/day GPG and '
        '~67 TJ/day large industrial (annual mean).'
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('GPG seasonality (empirical): ').bold = True
    p.add_run('winter-peaking, with eastern-states demand peaking in June at ~1.5x the annual '
              'mean. Winter/summer ratios are state-specific - Victoria 3.6x, NSW 2.4x, SA 1.9x, '
              'and Queensland ~1.0x (flat, dominated by baseload CCGTs such as Darling Downs).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Industrial seasonality: ').bold = True
    p.add_run('near-flat (winter/summer ~1.1x), consistent with baseload process load.')

    doc.add_heading('6.3 Three-tier curtailable demand stack', level=2)
    doc.add_paragraph(
        'GPG and large industrial are price-responsive in reality: when gas becomes expensive, '
        'gas generators switch to alternatives and some industrials curtail. To capture this '
        'without making demand fully endogenous, each tier is modelled as curtailable demand with '
        'a strike price - the nodal gas price above which the load sheds instead of being supplied. '
        'This avoids the unrealistic price/shortage overstatement that fully-inelastic demand would '
        'produce. The resulting merit order of load-shedding is:'
    )
    for tier, sp, note in [
        ('Gas-powered generation', '$22/GJ', 'sheds first - generators become uneconomic vs coal, batteries and imports'),
        ('Large industrial', '$120/GJ', 'sheds only in extreme scarcity - high cost of lost production, few alternatives'),
        ('Mass-market (firm)', '$300/GJ', 'value of lost load - residential/commercial, shed last via the shortage variable'),
    ]:
        q = doc.add_paragraph(style='List Bullet')
        q.add_run(f'{tier} ({sp}): ').bold = True
        q.add_run(note)
    doc.add_paragraph(
        'Formally, each tier adds its demand to the nodal balance and a non-negative curtailment '
        'variable bounded by that demand; the objective charges curtailment at the strike price '
        '(x1000 GJ/TJ). A tier sheds only when the marginal cost of supplying it would exceed its '
        'strike. GPG and industrial demand are held exogenous and flat across 2025-2050 (their '
        'own empirical seasonality is retained); they are not scaled by the winter or LNG demand '
        'multipliers, which apply to mass-market demand only.'
    )

    doc.add_heading('6.4 Limitations and next steps', level=2)
    doc.add_paragraph(
        'GPG/industrial levels are exogenous (price-responsive in dispatch via the strike, but not '
        'a function of the electricity market), and held flat in real terms to 2050 given genuine '
        'uncertainty about GPG\'s future role. A future extension would make GPG demand endogenous '
        'to electricity-market dispatch. The ~2 TJ/d Sydney residual overlap is immaterial and left '
        'unadjusted. Strike prices are tunable parameters (data/curtailment_params.csv).'
    )

    # Save
    doc.save('gas_market_model/Gas_Market_Model_Documentation.docx')
    print("Created Gas_Market_Model_Documentation.docx")

if __name__ == "__main__":
    create_docs()
